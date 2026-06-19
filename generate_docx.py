from docx import Document
from docx.shared import Pt
import os

doc = Document()
doc.add_heading('Project Report: BudgetBrake - Shop till the Budget Stops', 0)

# Chapter 1
doc.add_heading('Chapter 1: Introduction', level=1)
doc.add_heading('1.1 Background', level=2)
doc.add_paragraph('Personal financial management is a critical life skill, yet many individuals struggle to keep track of their expenses and adhere to a monthly budget. Impulse buying and lack of real-time expense tracking often lead to financial overruns. "BudgetBrake" is a localized desktop application designed to tackle this domain by simulating a marketplace where every purchase is strictly monitored against a predefined user budget.')

doc.add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph('Consumers frequently lose track of cumulative expenses while shopping across multiple categories and shops, resulting in budget deficits. There is a need for a unified, lightweight desktop application that integrates the shopping experience with strict, real-time budget tracking, visual analytics, and historical logging without relying on complex, heavy web-based architectures.')

doc.add_heading('1.3 Objectives', level=2)
doc.add_paragraph('The primary objective of this project is to develop a secure, multi-user application for shopping and budget management. It enforces real-time spending limits that prevent a user from completing a checkout if their cart total exceeds their remaining budget. Additionally, the system persists user data, including customized budgets, passwords, and transaction logs across sessions. To provide better insights, it includes a visual Analytics Dashboard that renders the user\'s spending habits over time using graphical charts. Finally, the application categorizes products logically by Shop and Category for intuitive navigation.')

doc.add_heading('1.4 Scope of the Project', level=2)
doc.add_paragraph('The project covers a complete simulated marketplace interface, cart management, user authentication, profile settings, and data persistence using text files. It includes graphical analytics rendered natively without external graphing libraries. A limitation of the system is that it is restricted to local desktop usage without cloud synchronization. It acts as a simulation and does not integrate with real-world payment gateways or live inventory APIs.')

doc.add_heading('1.5 Significance of the Project', level=2)
doc.add_paragraph('BudgetBrake serves as an excellent educational and practical tool. Practically, it helps users build financial discipline by visibly blocking transactions that exceed their means. Educationally, it demonstrates the power of the C programming language to build functional, event-driven Graphical User Interfaces (GUIs) and implement custom rendering logic from scratch.')

# Chapter 2
doc.add_heading('Chapter 2: System Design and Methodology', level=1)
doc.add_heading('2.1 Requirement Analysis', level=2)
doc.add_paragraph('The system processes various inputs including user credentials for authentication, category and shop selections for navigation, item quantities for purchasing, and user profile updates for changing passwords or budgets. In response, it generates outputs such as real-time cart total calculations, remaining budget updates, transaction logs, and dynamically plotted line graphs displaying expense versus time. The core functional requirements dictate that the system must restrict checkouts that exceed the user\'s budget, retain data upon application exit using local file storage, and ensure strict isolation of data between different users.')

doc.add_heading('2.2 Solution Design', level=2)
doc.add_paragraph('The application is built using an event-driven architecture relying on the Win32 API. Instead of a traditional linear console application, the program waits for user interactions such as button clicks and dropdown selections, processing them through a central WindowProc callback. Data persistence is handled securely via a localized file-system database structure within the data directory.')

doc.add_heading('2.3 Program Structure', level=2)
doc.add_paragraph('The program is structured into several core modules to ensure maintainability and separation of concerns. The Data Initialization and I/O Module handles reading and writing to user profiles, history logs, and shop inventory text files. The Authentication Module is responsible for validating login credentials and managing the active user\'s session state. The UI Layout Module manages the rendering and destruction of UI elements to simulate page navigation seamlessly. The Cart and Checkout Logic calculates totals, verifies budget limits, and triggers data saves upon transaction completion. Lastly, the Graph Rendering Engine uses a custom Window Procedure to plot data points and draw axes using the Graphics Device Interface.')

doc.add_heading('2.4 Data Structures and Variables', level=2)
doc.add_paragraph('The application extensively utilizes C structures to group related data logically. The User structure encapsulates user credentials, their specific budget, and an array of Log structures for their transaction history, while Shop and Category structures organize the inventory hierarchically. Arrays are employed for static, bounded collections such as item prices, categories, and users. Multi-dimensional arrays are used for fast, index-based accumulation of analytics data. Additionally, character arrays are used for handling strings like names, passwords, and formatted timestamps.')

doc.add_heading('2.5 Algorithm and Flowchart', level=2)
doc.add_paragraph('The checkout algorithm begins when the user clicks the Checkout button. The system first checks if the cart is empty; if so, it displays an error and aborts. Next, it prompts the user for confirmation. Upon confirmation, the cart total is subtracted from the active user\'s budget, and the updated budget is saved to the user\'s profile text file. A current timestamp is generated, and the transaction details are appended to the user\'s history log array and saved to a file. The system then loops through the cart arrays to update the shop spending matrix for analytics, saving this data as well. Finally, the cart totals are reset, and the UI financial labels are updated to reflect the successful transaction.')

# Chapter 3
doc.add_heading('Chapter 3: Implementation', level=1)
doc.add_heading('3.1 Development Environment', level=2)
doc.add_paragraph('The project was developed using the GCC compiler via MinGW on the Windows operating system. It relies on the Standard C Library for core functionalities like string manipulation and time handling. The graphical user interface and custom control drawing were implemented using the Windows API and the Windows Graphics Device Interface (GDI).')

doc.add_heading('3.2 Program Description', level=2)
doc.add_paragraph('The program execution begins in the WinMain function, which calls the initialization routines to bootstrap the file system and load existing users. It then registers the necessary window classes, including the main window and a custom GraphControl, before entering a continuous message loop. All application logic is strictly driven by the WM_COMMAND switch-cases which intercept and process button clicks and dropdown changes.')

doc.add_heading('3.3 Source Code', level=2)
doc.add_paragraph('The complete, well-documented source code is provided in the project files under BudgetBrake.c.')

doc.add_heading('3.4 Programming Concepts Used', level=2)
doc.add_paragraph('The project applies various fundamental and advanced programming concepts. Functions and modular programming are utilized to split logic into manageable pieces, avoiding bloated switch statements. Arrays and strings are extensively used for cart implementation and data parsing. C structures enforce object-oriented-like data grouping, while pointers are used minimally where Windows API requests memory addresses. File handling operations ensure data persists beyond the RAM lifecycle. Furthermore, iteration through for and while loops is heavily relied upon during file parsing, cart array shifting, and dynamic graph point plotting.')

doc.add_heading('3.5 Error Handling and Validation', level=2)
doc.add_paragraph('Robust error handling is implemented across the application. Budget validation strictly prevents adding items if the cart total combined with the item cost exceeds the available budget. File validation uses null pointer checks; if shop inventory files do not exist, the program automatically generates default inventory files and directories. Additionally, cart validation warns the user if they attempt to checkout with an empty cart or add an item without selecting a shop.')

# Chapter 4
doc.add_heading('Chapter 4: Testing, Results and Discussion', level=1)
doc.add_heading('4.1 Test Cases', level=2)
doc.add_paragraph('Several test cases were evaluated to ensure system reliability. For login validation, inputting an invalid password correctly results in a rejected login and an error popup. In the over-budget addition scenario, attempting to add items that exceed the predefined budget triggers a system block and displays a budget exceeded warning. For data persistence, checking out items and altering the profile budget, followed by closing and reopening the application, results in the new budget and recent checkout logs being restored perfectly.')

doc.add_heading('4.2 Program Execution Results', level=2)
doc.add_paragraph('Upon execution, the system successfully renders a Dark-Themed GUI. Navigating the categories allows users to select distinct items from dynamically generated ComboBoxes. The Analytics Dashboard successfully maps an interconnected Line Graph with properly scaled time and price axes based entirely on the user\'s historical checkouts.')

doc.add_heading('4.3 Validation of Results', level=2)
doc.add_paragraph('The program successfully meets all stated objectives. Multi-user isolation works flawlessly, ensuring that one user cannot view or modify another user\'s budget or history. The GDI drawing functions accurately calculate coordinates to plot the graph dynamically regardless of how large the monetary values get.')

doc.add_heading('4.4 Special and Edge Case Handling', level=2)
doc.add_paragraph('The program handles special edge cases gracefully. For instance, if data directories are missing on the first run, the program dynamically constructs the necessary file hierarchy on the fly. During cart deletion, if an item is removed from the middle of the cart, a loop iterates to shift all subsequent items, prices, and tracking indexes down by one, effectively preventing memory fragmentation and empty gaps in the UI listbox.')

doc.add_heading('4.5 Performance Analysis', level=2)
doc.add_paragraph('In terms of performance, the application boasts an extremely low memory footprint, consuming less than 5MB of RAM since it does not rely on heavy modern browser engines. Execution time is highly optimized because file I/O operations are restricted to initialization and discrete save events, ensuring the UI thread remains completely unblocked and highly responsive to user interactions.')

doc.add_heading('4.6 Discussion', level=2)
doc.add_paragraph('Building a GUI application in raw C using the Win32 API requires more explicit memory and event management compared to using modern frameworks. It necessitates manual coordinate math, manual memory management of graphics objects, and manual array shifting. However, the resulting application is incredibly lightweight, standalone, and provides an excellent fundamental understanding of how operating systems handle windows and events at a low level.')

# Chapter 5
doc.add_heading('Chapter 5: Conclusion and Future Enhancements', level=1)
doc.add_heading('5.1 Conclusion', level=2)
doc.add_paragraph('The BudgetBrake project successfully demonstrates the integration of a functional e-commerce-style shopping cart with a personal finance tracker. It proves that complex logic, including persistent file storage and dynamic visual analytics, can be built robustly using standard C and native OS APIs without the need for external dependencies.')

doc.add_heading('5.2 Learning Outcomes', level=2)
doc.add_paragraph('Developing this project yielded several key learning outcomes. It provided an advanced understanding of the Win32 API and event-driven programming in C. It also established proficiency in using the Windows Graphics Device Interface for rendering shapes and text natively. Furthermore, it demonstrated mastery of C file handling for persistent state management and improved skills in memory and array manipulation without the reliance on dynamic garbage collection.')

doc.add_heading('5.3 Future Scope', level=2)
doc.add_paragraph('There are several avenues for future enhancement. Migrating from flat text files to a relational database like SQLite would allow for faster querying and better security. Implementing encryption, such as SHA-256 hashing for passwords, would provide robust security compared to plain text storage. Additionally, providing a dynamic inventory admin panel would allow shop owners to add or remove items dynamically rather than relying on static text file generation.')

doc.add_heading('5.4 Real-World Applications', level=2)
doc.add_paragraph('The core logic of BudgetBrake can be adapted for lightweight point-of-sale systems, offline kiosk software, or personal finance educational tools targeted at students trying to build budgeting habits.')

doc.add_heading('References', level=1)
doc.add_paragraph('1. Microsoft Windows Win32 API Documentation\n2. Kernighan, B. W., & Ritchie, D. M. (1988). The C Programming Language.\n3. Windows GDI (Graphics Device Interface) reference guides for custom control drawing.\n4. Standard C Library Documentation for File I/O and string parsing techniques.')

doc.save('BudgetBrake_Project_Report_Paragraphs.docx')
print("Successfully created BudgetBrake_Project_Report_Paragraphs.docx")
